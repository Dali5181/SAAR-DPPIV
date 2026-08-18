"""
Formatted diagnostic report generation for DPP-IV inhibitor predictions.

Produces plain-text and Word (docx) reports from the structured dict
returned by ``RationalDesigner.generate_diagnostic_report()``.
"""
from __future__ import annotations

import os
from datetime import datetime


def generate_report_text(report_data: dict) -> str:
    """Render a structured report dict as formatted plain text."""
    lines: list[str] = []
    _hr = "=" * 72

    lines.append(_hr)
    lines.append("  AI-Driven DPP-IV Inhibitor Diagnostic Report")
    lines.append(f"  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(_hr)

    # --- Header ---
    seq = report_data["peptide_sequence"]
    lines.append(f"\nPeptide Sequence : {seq}")
    lines.append(f"Sequence Length  : {report_data['sequence_length']} residues")
    mw = report_data.get("molecular_weight_approx", 0)
    lines.append(f"Approx. MW       : {mw:.1f} Da")

    # --- Section 1: Activity Prediction ---
    lines.append(f"\n{'-' * 72}")
    lines.append("  Section 1: Activity Prediction")
    lines.append(f"{'-' * 72}")
    act = report_data["activity_prediction"]
    lines.append(f"  Predicted Class : {act['predicted_class']}")
    prob = act["class_probabilities"]
    lines.append(f"  P(active)       : {prob['active']:.4f}")
    lines.append(f"  P(inactive)     : {prob['inactive']:.4f}")
    lines.append(f"  Confidence      : {act['confidence']:.2%}")
    lines.append(f"  Predicted pIC50 : {act['predicted_pIC50']:.3f}")

    # --- Section 2: Binding Analysis ---
    lines.append(f"\n{'-' * 72}")
    lines.append("  Section 2: Binding Analysis")
    lines.append(f"{'-' * 72}")
    ba = report_data["binding_analysis"]

    lines.append("\n  Attention profile (normalised):")
    scores = ba["per_position_scores"]
    for i, s in enumerate(scores):
        bar = "#" * int(s * 30)
        residue = seq[i] if i < len(seq) else "?"
        marker = " *" if (i + 1) in ba["strong_positions"] else ""
        lines.append(f"    {residue}{i+1:>3d}  {s:.3f}  |{bar}{marker}")

    lines.append(f"\n  Strong binding positions : {ba['strong_positions']}")
    lines.append(f"  Weak binding positions   : {ba['weak_positions']}")

    if ba["key_binding_residues"]:
        lines.append("\n  Key binding residues:")
        for r in ba["key_binding_residues"]:
            lines.append(
                f"    {r['residue']}{r['position']}  "
                f"(score={r['attention_score']:.4f}, role={r['role']})"
            )

    # --- Section 3: Optimisation Suggestions ---
    lines.append(f"\n{'-' * 72}")
    lines.append("  Section 3: Optimisation Suggestions")
    lines.append(f"{'-' * 72}")

    for pocket_key in ("S1_pocket", "S2_pocket"):
        muts = report_data["mutation_suggestions"].get(pocket_key, [])
        label = pocket_key.replace("_", " ")
        lines.append(f"\n  {label} mutations (top {len(muts)}):")
        if not muts:
            lines.append("    (none suggested)")
        for m in muts:
            lines.append(
                f"    Pos {m['position']:>2d}: {m['original']} -> {m['suggested']}  "
                f"impact={m['expected_impact']:.3f}"
            )
            lines.append(f"           {m['rationale']}")

    mods = report_data.get("stability_modifications", [])
    lines.append(f"\n  Stability modifications ({len(mods)}):")
    for mod in mods:
        pos_str = ", ".join(str(p) for p in mod["positions"])
        lines.append(f"    [{mod['type']}] at position(s) {pos_str}")
        lines.append(f"      {mod['description']}")
        lines.append(f"      Benefit: {mod['expected_benefit']}")
        lines.append(f"      Feasibility: {mod['feasibility']}")

    # --- Section 4: Overall Assessment ---
    lines.append(f"\n{'-' * 72}")
    lines.append("  Section 4: Overall Assessment")
    lines.append(f"{'-' * 72}")
    oa = report_data["overall_assessment"]
    lines.append(f"  Score : {oa['score']:.1f} / {oa['max_score']:.0f}")
    lines.append(f"  Tier  : {oa['tier']}")
    lines.append(f"\n  {oa['summary']}")

    lines.append(f"\n{_hr}")
    lines.append("  End of Report")
    lines.append(_hr)

    return "\n".join(lines)


# ------------------------------------------------------------------
# Word document generation
# ------------------------------------------------------------------

def generate_report_docx(report_data: dict, output_path: str) -> str:
    """Create a Word document diagnostic report.

    Parameters
    ----------
    report_data : dict
        Structured report from ``RationalDesigner.generate_diagnostic_report()``.
    output_path : str
        File path for the ``.docx`` output.

    Returns
    -------
    str  The resolved output path.
    """
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # ---- Cover ----
    for _ in range(4):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_para.add_run("AI-Driven DPP-IV Inhibitor\nDiagnostic Report")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph()

    seq = report_data["peptide_sequence"]
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = sub_para.add_run(f"Peptide: {seq}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run(datetime.now().strftime("%Y-%m-%d"))
    run.font.size = Pt(12)

    doc.add_page_break()

    # ---- Section 1: Activity Prediction ----
    doc.add_heading("1  Activity Prediction", level=1)
    act = report_data["activity_prediction"]

    tbl = doc.add_table(rows=5, cols=2, style="Light List Accent 1")
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    _kv_rows = [
        ("Predicted Class", act["predicted_class"]),
        ("P(active)", f"{act['class_probabilities']['active']:.4f}"),
        ("P(inactive)", f"{act['class_probabilities']['inactive']:.4f}"),
        ("Confidence", f"{act['confidence']:.2%}"),
        ("Predicted pIC50", f"{act['predicted_pIC50']:.3f}"),
    ]
    for i, (k, v) in enumerate(_kv_rows):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = str(v)

    doc.add_paragraph()

    # ---- Section 2: Binding Analysis ----
    doc.add_heading("2  Binding Analysis", level=1)
    ba = report_data["binding_analysis"]

    doc.add_heading("2.1  Attention Profile", level=2)
    scores = ba["per_position_scores"]
    n_cols = min(len(scores), 20)
    if n_cols > 0:
        for chunk_start in range(0, len(scores), 20):
            chunk_end = min(chunk_start + 20, len(scores))
            chunk_len = chunk_end - chunk_start
            tbl = doc.add_table(rows=2, cols=chunk_len + 1, style="Table Grid")
            tbl.rows[0].cells[0].text = "Residue"
            tbl.rows[1].cells[0].text = "Score"
            for j, idx in enumerate(range(chunk_start, chunk_end)):
                residue = seq[idx] if idx < len(seq) else "?"
                tbl.rows[0].cells[j + 1].text = f"{residue}{idx+1}"
                tbl.rows[1].cells[j + 1].text = f"{scores[idx]:.2f}"
            doc.add_paragraph()

    doc.add_heading("2.2  Key Binding Residues", level=2)
    if ba["key_binding_residues"]:
        tbl = doc.add_table(
            rows=len(ba["key_binding_residues"]) + 1, cols=4,
            style="Light List Accent 1",
        )
        for ci, hdr in enumerate(["Position", "Residue", "Attention Score", "Role"]):
            tbl.rows[0].cells[ci].text = hdr
        for ri, r in enumerate(ba["key_binding_residues"], start=1):
            tbl.rows[ri].cells[0].text = str(r["position"])
            tbl.rows[ri].cells[1].text = r["residue"]
            tbl.rows[ri].cells[2].text = f"{r['attention_score']:.4f}"
            tbl.rows[ri].cells[3].text = r["role"]
    else:
        doc.add_paragraph("No strongly binding residues identified.")

    doc.add_paragraph()

    # ---- Section 3: Optimisation Suggestions ----
    doc.add_heading("3  Optimisation Suggestions", level=1)

    doc.add_heading("3.1  Pocket-Targeted Mutations", level=2)
    for pocket_key, label in [("S1_pocket", "S1 Pocket"),
                               ("S2_pocket", "S2 Pocket")]:
        muts = report_data["mutation_suggestions"].get(pocket_key, [])
        doc.add_heading(f"3.1.x  {label}", level=3)
        if not muts:
            doc.add_paragraph("No mutations suggested for this pocket.")
            continue
        tbl = doc.add_table(rows=len(muts) + 1, cols=5,
                            style="Light List Accent 1")
        for ci, hdr in enumerate(
                ["Pos", "Original", "Suggested", "Impact", "Rationale"]):
            tbl.rows[0].cells[ci].text = hdr
        for ri, m in enumerate(muts, start=1):
            tbl.rows[ri].cells[0].text = str(m["position"])
            tbl.rows[ri].cells[1].text = m["original"]
            tbl.rows[ri].cells[2].text = m["suggested"]
            tbl.rows[ri].cells[3].text = f"{m['expected_impact']:.3f}"
            tbl.rows[ri].cells[4].text = m["rationale"]
        doc.add_paragraph()

    doc.add_heading("3.2  Stability Modifications", level=2)
    mods = report_data.get("stability_modifications", [])
    if mods:
        tbl = doc.add_table(rows=len(mods) + 1, cols=5,
                            style="Light List Accent 1")
        for ci, hdr in enumerate(
                ["Type", "Positions", "Description", "Benefit", "Feasibility"]):
            tbl.rows[0].cells[ci].text = hdr
        for ri, mod in enumerate(mods, start=1):
            tbl.rows[ri].cells[0].text = mod["type"]
            tbl.rows[ri].cells[1].text = ", ".join(str(p) for p in mod["positions"])
            tbl.rows[ri].cells[2].text = mod["description"]
            tbl.rows[ri].cells[3].text = mod["expected_benefit"]
            tbl.rows[ri].cells[4].text = mod["feasibility"]
    else:
        doc.add_paragraph("No stability modifications identified.")

    doc.add_paragraph()

    # ---- Section 4: Overall Assessment ----
    doc.add_heading("4  Overall Assessment", level=1)
    oa = report_data["overall_assessment"]

    tbl = doc.add_table(rows=3, cols=2, style="Light List Accent 1")
    tbl.rows[0].cells[0].text = "Score"
    tbl.rows[0].cells[1].text = f"{oa['score']:.1f} / {oa['max_score']:.0f}"
    tbl.rows[1].cells[0].text = "Tier"
    tbl.rows[1].cells[1].text = oa["tier"]
    tbl.rows[2].cells[0].text = "Summary"
    tbl.rows[2].cells[1].text = oa["summary"]

    doc.add_paragraph()

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer_para.add_run(
        "Generated by AI-Driven DPP-IV Inhibitory Peptide Discovery Platform"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return os.path.abspath(output_path)
